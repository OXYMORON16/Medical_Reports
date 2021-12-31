import os
import docx
from tabulate import tabulate
import mysql.connector
mydb=mysql.connector.connect(host="localhost",password="123456",username="root",port=3307,database="project")
mycur=mydb.cursor()
def date():
    query_fetchdate="SELECT curdate()"
    mycur.execute(query_fetchdate)
    date=str(mycur.fetchall()[0][0])
    return date
def insert_reception():
    name=input("Please enter patient's name: ")
    age=int(input("Please enter patient's age: "))
    gender=input("Please enter patient's gender: ")
    address=input("Please enter patient's address: ")
    phone=input("Please enter patient's phone number: ")
    query_insert_reception="INSERT INTO reception VALUES('{}',{},'{}','{}','{}','{}')".format(name,age,gender,address,date(),phone)
    mycur.execute(query_insert_reception)
    mydb.commit()
    return "Data added successfully"
def insert_register(register_id):
    name=input("Please enter your name: ")
    age=int(input("Please enter your age: "))
    password=input("Please enter your password: ")
    dob=input("Please enter your date of birth in YYYY-MM-DD format: ")
    address=input("Please enter your address: ")
    mobile=input("Please enter your mobile number: ")
    query_insert_register="INSERT INTO register VALUES('{}','{}','{}','{}','{}','{}','{}')".format(register_id,name,age,password,dob,address,mobile)
    query_insert_login="INSERT INTO login VALUES('{}','{}')".format(register_id,password)
    mycur.execute(query_insert_register)
    mycur.execute(query_insert_login)
    mydb.commit()
    return "Registration Successful"
def see_all_patient_details():
    query_all_patient_details="SELECT * FROM reception"
    mycur.execute(query_all_patient_details)
    details=mycur.fetchall()
    return details
def see_one_patient_details():
    name=input("Please enter patient's name: ")
    query_one_patient_details="SELECT * FROM reception WHERE name='{}'".format(name)
    mycur.execute(query_one_patient_details)
    one_details=mycur.fetchall()
    return one_details
def output_all():
    details=see_all_patient_details()
    if details==[]:
        print("Please create a prescription form first")
    else:
        header = ['Name','Age','Gender','Address','Date','Phone']
        print(tabulate(details, header))
def output_one():
    details=see_one_patient_details()
    if details==[]:
        print("Please create a prescription form first")
    else:
        header = ['Name','Age','Gender','Address','Date','Phone']
        print(tabulate(details, header))
def output_all_word():
    doc=docx.Document()
    doc.add_heading("Vindhyachal Hospital",0)
    records=see_all_patient_details()
    menuTable=doc.add_table(rows=1,cols=6)
    menuTable.style='Table Grid'
    hdr_Cells = menuTable.rows[0].cells
    hdr_Cells[0].text = 'Name'
    hdr_Cells[1].text = 'Age'
    hdr_Cells[2].text = 'Gender'
    hdr_Cells[3].text = 'Address'
    hdr_Cells[4].text = 'Date'
    hdr_Cells[5].text = 'Phone'
    for name,age,gender,address,date,phone in records:
        row_Cells = menuTable.add_row().cells
        row_Cells[0].text = name
        row_Cells[1].text = str(age)
        row_Cells[2].text = gender
        row_Cells[3].text = address
        row_Cells[4].text = str(date)
        row_Cells[5].text = phone
    doc.save('output_all_patients.docx')
    os.system("start output_all_patients.docx")
    return ""
def output_one_word():
    doc=docx.Document()
    doc.add_heading("Vindhyachal Hospital",0)
    records=see_one_patient_details()
    menuTable=doc.add_table(rows=1,cols=6)
    menuTable.style='Table Grid'
    hdr_Cells = menuTable.rows[0].cells
    hdr_Cells[0].text = 'Name'
    hdr_Cells[1].text = 'Age'
    hdr_Cells[2].text = 'Gender'
    hdr_Cells[3].text = 'Address'
    hdr_Cells[4].text = 'Date'
    hdr_Cells[5].text = 'Phone'
    for name,age,gender,address,date,phone in records:
        row_Cells = menuTable.add_row().cells
        row_Cells[0].text = name
        row_Cells[1].text = str(age)
        row_Cells[2].text = gender
        row_Cells[3].text = address
        row_Cells[4].text = str(date)
        row_Cells[5].text = phone
    doc.save('output_one_patients.docx')
    os.system("start output_one_patients.docx")
def delete_all_patients():
    query_delete_all_patient="DELETE FROM reception"
    mycur.execute(query_delete_all_patient)
    mydb.commit()
    return "Patients records successfully deleted"
def delete_one_patient():
    name=input("Enter name of the patient to delete: ")
    query_delete_one_patient="DELETE FROM reception WHERE Name = '{}'".format(name)
    mycur.execute(query_delete_one_patient)
    mydb.commit()
    return "{}\'s record deleted successfully".format(name)
